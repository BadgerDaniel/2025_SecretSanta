import streamlit as st
import random
import secrets
import time
from io import BytesIO
from docx import Document

# ===========================
# Data
# ===========================

names = ["Ariel", "Kurt", "Scott", "Linda", "Daniel"]

# Directional only — reverse is allowed
# Use list of tuples to handle cases where one person may have multiple blocked recipients
last_year_gifts = [
    ("Kurt", "Ariel"),
    ("Ariel", "Kurt"),
    ("Scott", "Ariel"),
    ("Linda", "Scott"),
    ("Daniel", "Linda"),
    ("Ariel", "Daniel")
]

blocked_pairs = set(last_year_gifts)

# ===========================
# Pairing Generator
# ===========================

def generate_valid_pairs():
    # Use cryptographically secure random seed for true randomness
    # Combine time and secrets for maximum randomness
    seed = int(time.time() * 1000000) + secrets.randbits(32)
    random.seed(seed)
    
    for attempt in range(10000):
        shuffled = names.copy()
        random.shuffle(shuffled)
        pairs = dict(zip(names, shuffled))
        
        # Validate: no self-gifting and no blocked pairs
        is_valid = True
        for giver, receiver in pairs.items():
            if giver == receiver:
                is_valid = False
                break
            if (giver, receiver) in blocked_pairs:
                is_valid = False
                break
        
        if is_valid:
            return pairs
    
    return None

# ===========================
# DOCX Export
# ===========================

def generate_docx(pairs):
    doc = Document()
    doc.add_heading("Secret Santa Assignments", level=1)
    for giver, receiver in pairs.items():
        doc.add_paragraph(f"{giver} → {receiver}")
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ===========================
# Session State Setup
# ===========================

if "pairs" not in st.session_state:
    st.session_state.pairs = generate_valid_pairs()
    st.session_state.step = 0
    st.session_state.revealed = False
    st.session_state.initialized = True

pairs = st.session_state.pairs

# ===========================
# UI
# ===========================

st.title("🎅 Secret Santa (Sequential Reveal Version)")

# Add reset button in sidebar for generating new pairings
with st.sidebar:
    st.header("Options")
    if st.button("🔄 Generate New Pairings", help="This will reset and generate completely new random pairings"):
        # Clear session state to force regeneration
        for key in list(st.session_state.keys()):
            del st.session_state[key]
        st.rerun()
    
    # Debug: Show blocked pairs
    with st.expander("🔒 Blocked Pairings (Last Year)"):
        if blocked_pairs:
            for giver, receiver in sorted(blocked_pairs):
                st.text(f"{giver} → {receiver}")
        else:
            st.text("No blocked pairings")

if pairs is None:
    st.error("Could not generate a valid Secret Santa pairing.")
    st.stop()

step = st.session_state.step

# ===========================
# When finished
# ===========================

if step >= len(names):
    st.success("🎉 Everyone has seen their person! Happy gifting!")
    # Export button
    docx_data = generate_docx(pairs)
    st.download_button(
        label="📄 Download Word Doc of All Pairings",
        data=docx_data,
        file_name="secret_santa_assignments.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    st.stop()

# ===========================
# Active Person View
# ===========================

current_person = names[step]

st.header(f"👤 {current_person}, it's your turn!")

if not st.session_state.revealed:
    if st.button("Reveal who you're gifting to"):
        st.session_state.revealed = True
else:
    recipient = pairs[current_person]
    st.success(f"🎁 You are gifting to: **{recipient}**")
    
    # Show who to pass to next (if any left)
    if step + 1 < len(names):
        next_person = names[step + 1]
        st.info(f"👉 When ready, pass the device to **{next_person}**.")
    else:
        st.info("🏁 You're the last person!")
    
    if st.button("Next person (hide this)"):
        st.session_state.revealed = False
        st.session_state.step += 1

